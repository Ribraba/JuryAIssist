"""
Fenêtre principale de l'application Flet - Style minimaliste Apple-like.

Assemble tous les composants et gère la logique métier:
- Lecture audio via VLCAudioPlayer
- Transcription via WhisperTranscriber
- Contrôle pédale Olympus (optionnel)
- Export TXT/DOCX
"""
import flet as ft
from pathlib import Path
from typing import Optional, List
import threading

from src.gui_flet.theme import AppColors, AppSpacing, get_theme
from src.gui_flet.components.sidebar import Sidebar
from src.gui_flet.components.audio_player import AudioPlayer
from src.gui_flet.components.editor_panel import EditorPanel
from src.gui_flet.components.transcription_view import TranscriptionView

# Modules métier
from src.audio.controller import AudioController
from src.audio.vlc_player import VLCAudioPlayer
from src.audio.player import PlayerState
from src.transcription.whisper_transcriber import WhisperTranscriber
from src.transcription.transcriber import TranscriptionSegment
from src.devices.olympus_pedal import OlympusPedal
from src.devices.pedal import PedalAction


class MainWindow:
    """
    Fenêtre principale de l'application.

    Coordonne tous les composants et la logique métier.
    """

    def __init__(self, page: ft.Page):
        """
        Initialise la fenêtre principale.

        Args:
            page: Page Flet
        """
        self.page = page

        # Configuration de la page
        self._setup_page()

        # État
        self.current_audio_file: Optional[str] = None
        self.current_transcript: List[TranscriptionSegment] = []

        # Audio
        self.player = VLCAudioPlayer()
        self.controller = AudioController(self.player)

        # Pédale (optionnelle)
        self.pedal: Optional[OlympusPedal] = None
        self._init_pedal()

        # Composants UI
        self.sidebar = Sidebar(
            on_import_clicked=self._on_import_clicked,
            on_settings_clicked=self._on_settings_clicked,
            on_file_selected=self._on_file_selected,
        )

        self.audio_player = AudioPlayer(
            on_play=self._on_play,
            on_pause=self._on_pause,
            on_stop=self._on_stop,
            on_skip_forward=self._on_skip_forward,
            on_skip_backward=self._on_skip_backward,
            on_seek=self._on_seek,
            on_speed_change=self._on_speed_change,
            on_volume_change=self._on_volume_change,
        )

        self.editor_panel = EditorPanel(
            on_export_txt=self._on_export_txt,
            on_export_docx=self._on_export_docx,
        )

        self.transcription_view = TranscriptionView(
            on_segment_click=self._on_segment_click,
        )

        # Connecter les signaux audio
        self._connect_signals()

        # Construire l'interface
        self._build_ui()

        # Timer pour mettre à jour la position
        self._position_timer = None
        self._start_position_timer()

    def _setup_page(self):
        """Configure la page Flet."""
        self.page.title = "JuryAIssist - Transcription Audio"
        self.page.theme = get_theme()
        self.page.bgcolor = AppColors.BACKGROUND
        self.page.padding = 0
        self.page.window_width = 1440
        self.page.window_height = 960
        self.page.window_min_width = 1200
        self.page.window_min_height = 800

    def _build_ui(self):
        """Construit l'interface utilisateur."""
        # Layout principal: Sidebar | Content
        content_area = ft.Column(
            [
                # En-tête avec badge pédale
                self._build_header(),

                ft.Container(height=AppSpacing.MD),

                # Éditeur (prend tout l'espace)
                self.editor_panel,

                ft.Container(height=AppSpacing.MD),

                # Vue transcription
                self.transcription_view,

                ft.Container(height=AppSpacing.MD),

                # Lecteur audio
                self.audio_player,
            ],
            spacing=0,
            expand=True,
        )

        main_row = ft.Row(
            [
                self.sidebar,
                ft.Container(
                    content=content_area,
                    padding=AppSpacing.LG,
                    expand=True,
                ),
            ],
            spacing=0,
            expand=True,
        )

        self.page.add(main_row)
        self.page.update()

    def _build_header(self) -> ft.Container:
        """Construit l'en-tête avec le badge de pédale."""
        pedal_status = "Pédale connectée" if self.pedal and self.pedal.is_connected() else "Pédale non détectée"
        pedal_icon = "check_circle" if self.pedal and self.pedal.is_connected() else "circle"
        pedal_color = AppColors.SUCCESS if self.pedal and self.pedal.is_connected() else AppColors.TEXT_TERTIARY

        return ft.Container(
            content=ft.Row(
                [
                    ft.Icon(pedal_icon, size=16, color=pedal_color),
                    ft.Text(
                        pedal_status,
                        size=12,
                        color=pedal_color,
                    ),
                ],
                spacing=AppSpacing.XS,
                alignment=ft.MainAxisAlignment.END,
            ),
        )

    def _connect_signals(self):
        """Connecte les signaux du contrôleur audio."""
        self.controller.position_changed.connect(self._on_position_changed)
        self.controller.duration_changed.connect(self._on_duration_changed)
        self.controller.state_changed.connect(self._on_state_changed)

    def _start_position_timer(self):
        """Démarre le timer de mise à jour de la position."""
        def update_position():
            if self.controller:
                position = self.controller.get_position()
                self._on_position_changed(position)

        # Flet ne supporte pas les timers Qt, on utilise un threading.Timer
        def timer_loop():
            while True:
                update_position()
                threading.Event().wait(0.1)  # 100ms

        timer_thread = threading.Thread(target=timer_loop, daemon=True)
        timer_thread.start()

    # === Callbacks UI ===

    def _on_import_clicked(self):
        """Gère le clic sur Importer."""
        def on_result(e: ft.FilePickerResultEvent):
            if e.files:
                file_path = e.files[0].path
                self._load_audio_file(file_path)

        file_picker = ft.FilePicker(on_result=on_result)
        self.page.overlay.append(file_picker)
        self.page.update()
        file_picker.pick_files(
            allowed_extensions=["mp3", "wav", "m4a", "flac", "ogg", "dss"],
            dialog_title="Ouvrir un fichier audio",
        )

    def _on_settings_clicked(self):
        """Gère le clic sur Paramètres."""
        # TODO: Implémenter un dialogue de paramètres
        self._show_dialog(
            "Paramètres",
            "Les paramètres seront disponibles prochainement.",
        )

    def _on_file_selected(self, filename: str):
        """
        Gère la sélection d'un fichier.

        Args:
            filename: Nom du fichier
        """
        # TODO: Gérer la sélection de fichier
        pass

    def _on_play(self):
        """Lance la lecture."""
        self.controller.play()

    def _on_pause(self):
        """Met en pause."""
        self.controller.pause()

    def _on_stop(self):
        """Arrête la lecture."""
        self.controller.stop()

    def _on_skip_forward(self):
        """Avance de 5 secondes."""
        self.controller.skip_forward(5.0)

    def _on_skip_backward(self):
        """Recule de 5 secondes."""
        self.controller.skip_backward(5.0)

    def _on_seek(self, position: float):
        """
        Change la position de lecture.

        Args:
            position: Nouvelle position en secondes
        """
        self.controller.seek(position)

    def _on_speed_change(self, speed: float):
        """
        Change la vitesse de lecture.

        Args:
            speed: Nouvelle vitesse (ex: 1.5)
        """
        self.controller.set_speed(speed)

    def _on_volume_change(self, volume: int):
        """
        Change le volume.

        Args:
            volume: Nouveau volume (0-100)
        """
        self.controller.set_volume(volume)

    def _on_segment_click(self, timestamp: float):
        """
        Gère le clic sur un segment de transcription.

        Args:
            timestamp: Position en secondes
        """
        self.controller.seek(timestamp)

    def _on_export_txt(self, text: str):
        """
        Exporte en TXT.

        Args:
            text: Texte à exporter
        """
        if not text:
            self._show_dialog("Attention", "Aucun texte à exporter.")
            return

        def on_result(e: ft.FilePickerResultEvent):
            if e.path:
                try:
                    with open(e.path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    self._show_dialog("Succès", f"Fichier exporté: {Path(e.path).name}")
                except Exception as ex:
                    self._show_dialog("Erreur", f"Erreur lors de l'export:\n{ex}")

        file_picker = ft.FilePicker(on_result=on_result)
        self.page.overlay.append(file_picker)
        self.page.update()
        file_picker.save_file(
            file_name="transcription.txt",
            allowed_extensions=["txt"],
            dialog_title="Exporter en TXT",
        )

    def _on_export_docx(self, text: str):
        """
        Exporte en DOCX.

        Args:
            text: Texte à exporter
        """
        if not text:
            self._show_dialog("Attention", "Aucun texte à exporter.")
            return

        def on_result(e: ft.FilePickerResultEvent):
            if e.path:
                try:
                    from docx import Document

                    doc = Document()
                    doc.add_heading('Transcription Audio', 0)

                    paragraphs = text.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para)

                    doc.save(e.path)
                    self._show_dialog("Succès", f"Fichier exporté: {Path(e.path).name}")
                except ImportError:
                    self._show_dialog(
                        "Module manquant",
                        "Le module python-docx n'est pas installé.\n"
                        "Installez-le avec: pip install python-docx"
                    )
                except Exception as ex:
                    self._show_dialog("Erreur", f"Erreur lors de l'export:\n{ex}")

        file_picker = ft.FilePicker(on_result=on_result)
        self.page.overlay.append(file_picker)
        self.page.update()
        file_picker.save_file(
            file_name="transcription.docx",
            allowed_extensions=["docx"],
            dialog_title="Exporter en DOCX",
        )

    # === Signaux Audio ===

    def _on_position_changed(self, position: float):
        """
        Appelé quand la position change.

        Args:
            position: Position en secondes
        """
        self.audio_player.set_position(position)
        self.transcription_view.set_position(position)

    def _on_duration_changed(self, duration: float):
        """
        Appelé quand la durée change.

        Args:
            duration: Durée en secondes
        """
        self.audio_player.set_duration(duration)

    def _on_state_changed(self, state: PlayerState):
        """
        Appelé quand l'état change.

        Args:
            state: Nouvel état
        """
        is_playing = (state == PlayerState.PLAYING)
        self.audio_player.set_playing_state(is_playing)

    # === Logique métier ===

    def _load_audio_file(self, file_path: str):
        """
        Charge un fichier audio.

        Args:
            file_path: Chemin du fichier
        """
        if self.controller.load_file(file_path):
            self.current_audio_file = file_path
            filename = Path(file_path).name

            # Mettre à jour l'interface
            self.editor_panel.set_filename(filename)
            self.sidebar.clear_files()
            self.sidebar.add_file(filename, selected=True)

            # Demander si on veut transcrire
            self._ask_transcription()
        else:
            self._show_dialog("Erreur", f"Impossible de charger le fichier:\n{file_path}")

    def _ask_transcription(self):
        """Demande si l'utilisateur veut transcrire."""
        def on_yes(e):
            dialog.open = False
            self.page.update()
            self._start_transcription()

        def on_no(e):
            dialog.open = False
            self.page.update()

        dialog = ft.AlertDialog(
            title=ft.Text("Transcription"),
            content=ft.Text("Voulez-vous transcrire ce fichier audio maintenant?"),
            actions=[
                ft.TextButton("Non", on_click=on_no),
                ft.TextButton("Oui", on_click=on_yes),
            ],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dialog
        dialog.open = True
        self.page.update()

    def _start_transcription(self):
        """Lance la transcription."""
        if not self.current_audio_file:
            self._show_dialog("Attention", "Aucun fichier audio chargé.")
            return

        # Afficher un dialogue de progression
        progress_bar = ft.ProgressBar(width=400)
        progress_text = ft.Text("Transcription en cours...")

        dialog = ft.AlertDialog(
            title=ft.Text("Transcription"),
            content=ft.Column(
                [progress_text, progress_bar],
                tight=True,
            ),
            modal=True,
        )

        self.page.dialog = dialog
        dialog.open = True
        self.page.update()

        # Lancer la transcription dans un thread
        def transcribe():
            try:
                transcriber = WhisperTranscriber(model_size="base")
                result = transcriber.transcribe(self.current_audio_file, language="fr")

                # Mettre à jour l'UI dans le thread principal
                def update_ui():
                    dialog.open = False
                    self.page.update()

                    if result.status.name == "COMPLETED":
                        self.current_transcript = result.segments
                        self._on_transcription_completed(result.segments)
                    else:
                        self._show_dialog("Erreur", f"Erreur: {result.error_message}")

                self.page.run_task(update_ui)

            except Exception as e:
                def show_error():
                    dialog.open = False
                    self.page.update()
                    self._show_dialog("Erreur", f"Erreur lors de la transcription:\n{e}")

                self.page.run_task(show_error)

        thread = threading.Thread(target=transcribe, daemon=True)
        thread.start()

    def _on_transcription_completed(self, segments: List[TranscriptionSegment]):
        """
        Appelé quand la transcription est terminée.

        Args:
            segments: Segments de transcription
        """
        # Mettre à jour la vue de transcription
        self.transcription_view.set_segments(segments)

        # Mettre à jour l'éditeur
        text_lines = [seg.text for seg in segments]
        self.editor_panel.set_text("\n\n".join(text_lines))

        self._show_dialog(
            "Succès",
            f"Transcription terminée!\n\n{len(segments)} segments créés."
        )

    def _init_pedal(self):
        """Initialise la pédale (silencieux si absente)."""
        try:
            self.pedal = OlympusPedal()
            if self.pedal.detect() and self.pedal.connect():
                self.pedal.action_triggered.connect(self._on_pedal_action)
                print("✅ Pédale connectée")
        except Exception:
            pass

    def _on_pedal_action(self, action: PedalAction):
        """
        Gère les actions de la pédale.

        Args:
            action: Action de la pédale
        """
        if action == PedalAction.PLAY_PAUSE:
            self.controller.toggle_play_pause()
        elif action == PedalAction.SKIP_FORWARD:
            self.controller.skip_forward(5.0)
        elif action == PedalAction.SKIP_BACKWARD:
            self.controller.skip_backward(5.0)
        elif action == PedalAction.STOP:
            self.controller.stop()

    def _show_dialog(self, title: str, message: str):
        """
        Affiche un dialogue simple.

        Args:
            title: Titre du dialogue
            message: Message
        """
        def on_close(e):
            dialog.open = False
            self.page.update()

        dialog = ft.AlertDialog(
            title=ft.Text(title),
            content=ft.Text(message),
            actions=[ft.TextButton("OK", on_click=on_close)],
            actions_alignment=ft.MainAxisAlignment.END,
        )

        self.page.dialog = dialog
        dialog.open = True
        self.page.update()

    def cleanup(self):
        """Nettoyage avant fermeture."""
        if self.controller:
            self.controller.release()
        if self.pedal:
            self.pedal.disconnect()
