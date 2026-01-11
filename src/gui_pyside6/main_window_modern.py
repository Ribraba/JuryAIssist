"""
Fenêtre principale moderne de JuryAIssist avec PySide6.
Design inspiré macOS avec sidebar et éditeur unique.
"""

from pathlib import Path
from typing import Optional, List
import threading

from PySide6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QSlider, QTextEdit, QFileDialog,
    QMessageBox, QProgressDialog, QFrame, QListWidget,
    QListWidgetItem, QComboBox, QSpinBox, QDialog,
    QDialogButtonBox, QSizePolicy, QApplication, QMenuBar, QMenu
)
from PySide6.QtCore import Qt, QTimer, Slot, QSize, QPropertyAnimation, QEasingCurve, Property, QEvent, Signal, QObject
from PySide6.QtGui import QIcon, QFont, QAction, QTextCursor, QTextCharFormat, QColor, QKeySequence, QKeyEvent, QShortcut

from src.audio.controller import AudioController
from src.audio.vlc_player import VLCAudioPlayer
from src.audio.player import PlayerState
from src.transcription.whisper_transcriber import WhisperTranscriber
from src.devices.olympus_pedal import OlympusPedal
from .styles import MAIN_STYLE, COLORS
from .icon_loader import IconLoader
from .modern_progress import ModernProgressDialog


class TranscriptionItem:
    """Item de transcription."""
    def __init__(self, file_path: Path, name: str):
        self.file_path = file_path
        self.name = name
        self.transcription_text = ""
        self.segments = []


class AnimatedButton(QPushButton):
    """Bouton avec animation de scale au hover."""

    def __init__(self, text: str, parent=None):
        super().__init__(text, parent)
        self._scale = 1.0

    def enterEvent(self, event):
        """Animation au survol."""
        super().enterEvent(event)
        self.animate_scale(1.05)

    def leaveEvent(self, event):
        """Animation à la sortie."""
        super().leaveEvent(event)
        self.animate_scale(1.0)

    def animate_scale(self, target_scale: float):
        """Anime l'échelle du bouton."""
        # Note: PySide6 ne supporte pas directement le scale via QPropertyAnimation sur les widgets
        # On simule avec la font size pour l'instant
        pass

    def get_scale(self):
        return self._scale

    def set_scale(self, scale):
        self._scale = scale

    scale = Property(float, get_scale, set_scale)


class SettingsDialog(QDialog):
    """Dialogue de paramètres moderne."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Paramètres")
        self.setMinimumWidth(450)
        self.setStyleSheet(MAIN_STYLE)

        layout = QVBoxLayout()
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)

        # Titre
        title = QLabel("Paramètres de transcription")
        title.setObjectName("header_label")
        layout.addWidget(title)

        # Modèle Whisper
        model_label = QLabel("Modèle Whisper")
        model_label.setObjectName("secondary_label")
        layout.addWidget(model_label)

        self.model_combo = QComboBox()
        self.model_combo.addItems(["tiny", "base", "small", "medium", "large"])
        self.model_combo.setCurrentText("base")
        layout.addWidget(self.model_combo)

        layout.addSpacing(8)

        # Langue
        lang_label = QLabel("Langue de transcription")
        lang_label.setObjectName("secondary_label")
        layout.addWidget(lang_label)

        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["Français", "English", "Español", "Deutsch", "Italiano"])
        self.lang_combo.setCurrentIndex(0)
        layout.addWidget(self.lang_combo)

        layout.addSpacing(8)

        # Durée du skip
        skip_label = QLabel("Durée du skip (secondes)")
        skip_label.setObjectName("secondary_label")
        layout.addWidget(skip_label)

        self.skip_spin = QSpinBox()
        self.skip_spin.setRange(1, 30)
        self.skip_spin.setValue(5)
        layout.addWidget(self.skip_spin)

        layout.addStretch()

        # Boutons
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self.setLayout(layout)

    def get_language_code(self):
        """Convertit le nom de langue en code."""
        lang_map = {
            "Français": "fr",
            "English": "en",
            "Español": "es",
            "Deutsch": "de",
            "Italiano": "it"
        }
        return lang_map.get(self.lang_combo.currentText(), "fr")


class TranscriptionSignals(QObject):
    """Signaux pour communication thread-safe entre worker et GUI."""
    completed = Signal(object, object)  # (result, progress_dialog)
    error = Signal(str, object)  # (error_message, progress_dialog)


class ModernMainWindow(QMainWindow):
    """Fenêtre principale moderne de JuryAIssist."""

    def __init__(self):
        super().__init__()

        # État de l'application
        self.transcriptions: List[TranscriptionItem] = []
        self.current_transcription: Optional[TranscriptionItem] = None
        self.is_modified = False
        self.skip_duration = 5.0
        self.whisper_model = "base"
        self.whisper_language = "fr"

        # Composants métier
        self.player = VLCAudioPlayer()
        self.controller = AudioController(self.player)
        self.transcriber = WhisperTranscriber()
        self.pedal = OlympusPedal()

        # Connecter les signaux
        self.controller.position_changed.connect(self._on_position_changed)
        self.controller.duration_changed.connect(self._on_duration_changed)
        self.controller.state_changed.connect(self._on_state_changed)
        self.controller.source_loaded.connect(self._on_source_loaded)

        self.pedal.action_triggered.connect(self._on_pedal_action)
        self.pedal.connected.connect(self._on_pedal_connected)
        self.pedal.disconnected.connect(self._on_pedal_disconnected)

        # Initialiser l'UI
        self._setup_ui()
        self._setup_menu_bar()
        self._apply_styles()
        self._setup_shortcuts()

        # Connecter la pédale
        self.pedal.connect()

    def _setup_ui(self):
        """Configure l'interface utilisateur."""
        self.setWindowTitle("JuryAIssist")
        self.setMinimumSize(1400, 900)

        # Widget central
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # Layout principal (horizontal: sidebar + contenu)
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # === SIDEBAR GAUCHE ===
        self.sidebar = self._create_sidebar()
        main_layout.addWidget(self.sidebar)

        # === ZONE PRINCIPALE ===
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setSpacing(0)
        content_layout.setContentsMargins(0, 0, 0, 0)

        # En-tête avec titre et badge pédale
        header = self._create_header()
        content_layout.addWidget(header)

        # Éditeur principal (prend tout l'espace)
        self.text_editor = QTextEdit()
        self.text_editor.setPlaceholderText("Sélectionnez ou importez un fichier audio pour commencer la transcription...")
        self.text_editor.textChanged.connect(self._on_text_changed)
        content_layout.addWidget(self.text_editor, 1)

        # Timeline en bas
        timeline_frame = self._create_timeline()
        content_layout.addWidget(timeline_frame)

        main_layout.addWidget(content_widget, 1)

        # Barre de statut
        self.statusBar().showMessage("Prêt")

    def _create_sidebar(self) -> QWidget:
        """Crée la sidebar gauche."""
        sidebar = QFrame()
        sidebar.setObjectName("sidebar")
        sidebar.setMinimumWidth(280)
        sidebar.setMaximumWidth(320)

        layout = QVBoxLayout(sidebar)
        layout.setSpacing(16)
        layout.setContentsMargins(16, 20, 16, 20)

        # Titre de l'app
        title = QLabel("JuryAIssist")
        title.setObjectName("sidebar_title")
        layout.addWidget(title)

        # Bouton Nouvel import
        self.btn_import = QPushButton(" Nouvel import")
        self.btn_import.setObjectName("white_button")
        self.btn_import.setIcon(IconLoader.get_icon("dossier", 20))
        self.btn_import.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_import.clicked.connect(self._on_import_file)
        self.btn_import.setMinimumHeight(44)
        layout.addWidget(self.btn_import)

        # Séparateur
        layout.addSpacing(8)

        # Label Transcriptions
        trans_label = QLabel("Transcriptions")
        trans_label.setObjectName("secondary_label")
        layout.addWidget(trans_label)

        # Liste des transcriptions
        self.transcriptions_list = QListWidget()
        self.transcriptions_list.itemClicked.connect(self._on_transcription_selected)
        layout.addWidget(self.transcriptions_list, 1)

        # Bouton Paramètres en bas
        layout.addStretch()
        self.btn_settings = QPushButton(" Paramètres")
        self.btn_settings.setIcon(IconLoader.get_icon("engrenage", 18))
        self.btn_settings.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_settings.clicked.connect(self._on_settings)
        layout.addWidget(self.btn_settings)

        return sidebar

    def _create_header(self) -> QWidget:
        """Crée l'en-tête avec titre et badge pédale."""
        header = QFrame()
        header.setStyleSheet(f"background-color: {COLORS['bg_primary']}; border-bottom: 1px solid {COLORS['border_light']};")
        header.setMinimumHeight(90)

        # Layout vertical principal
        layout = QVBoxLayout(header)
        layout.setContentsMargins(24, 16, 24, 16)
        layout.setSpacing(12)

        # Ligne du haut: Bouton "En savoir plus" + Badge pédale
        top_row = QHBoxLayout()
        top_row.setSpacing(16)

        # Bouton "En savoir plus" (NOIR avec texte blanc)
        self.btn_info = QPushButton("En savoir plus")
        self.btn_info.setObjectName("black_button")
        self.btn_info.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_info.clicked.connect(self._on_show_info)
        self.btn_info.setMinimumWidth(140)
        self.btn_info.setFixedHeight(32)
        top_row.addWidget(self.btn_info)

        top_row.addStretch()

        # Badge pédale (à droite)
        self.pedal_badge = QLabel("○ Pédale: Non connectée")
        self.pedal_badge.setObjectName("pedal_badge")
        self.pedal_badge.setProperty("connected", False)
        top_row.addWidget(self.pedal_badge)

        layout.addLayout(top_row)

        # Ligne du bas: Titre du fichier (AGRANDI)
        self.file_title = QLabel("Aucun fichier chargé")
        self.file_title.setObjectName("title_large")
        self.file_title.setWordWrap(True)
        layout.addWidget(self.file_title)

        return header

    def _create_timeline(self) -> QWidget:
        """Crée la timeline avec contrôles audio."""
        timeline = QFrame()
        timeline.setObjectName("timeline_frame")
        timeline.setMinimumHeight(140)

        layout = QVBoxLayout(timeline)
        layout.setSpacing(12)
        layout.setContentsMargins(24, 16, 24, 16)

        # Barre de progression
        progress_layout = QHBoxLayout()
        progress_layout.setSpacing(12)

        self.time_label = QLabel("0:00")
        self.time_label.setObjectName("secondary_label")
        self.time_label.setMinimumWidth(50)
        progress_layout.addWidget(self.time_label)

        self.progress_slider = QSlider(Qt.Orientation.Horizontal)
        self.progress_slider.setRange(0, 1000)
        self.progress_slider.setCursor(Qt.CursorShape.PointingHandCursor)
        self.progress_slider.sliderMoved.connect(self._on_slider_moved)
        progress_layout.addWidget(self.progress_slider, 1)

        self.duration_label = QLabel("0:00")
        self.duration_label.setObjectName("secondary_label")
        self.duration_label.setMinimumWidth(50)
        self.duration_label.setAlignment(Qt.AlignmentFlag.AlignRight)
        progress_layout.addWidget(self.duration_label)

        layout.addLayout(progress_layout)

        # Contrôles audio
        controls_layout = QHBoxLayout()
        controls_layout.setSpacing(12)

        # Boutons de contrôle avec icônes SVG
        self.btn_skip_back = QPushButton()
        self.btn_skip_back.setIcon(IconLoader.get_icon("skip_backward", 24))
        self.btn_skip_back.setIconSize(QSize(24, 24))
        self.btn_skip_back.setMinimumSize(48, 48)
        self.btn_skip_back.setObjectName("white_button")
        self.btn_skip_back.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_skip_back.clicked.connect(self._on_skip_back)
        controls_layout.addWidget(self.btn_skip_back)

        self.btn_play = QPushButton()
        self.btn_play.setIcon(IconLoader.get_icon("play", 28))
        self.btn_play.setIconSize(QSize(28, 28))
        self.btn_play.setMinimumSize(56, 56)
        self.btn_play.setObjectName("white_button")
        self.btn_play.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_play.clicked.connect(self._on_play)
        controls_layout.addWidget(self.btn_play)

        self.btn_skip_forward = QPushButton()
        self.btn_skip_forward.setIcon(IconLoader.get_icon("skip_forward", 24))
        self.btn_skip_forward.setIconSize(QSize(24, 24))
        self.btn_skip_forward.setMinimumSize(48, 48)
        self.btn_skip_forward.setObjectName("white_button")
        self.btn_skip_forward.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_skip_forward.clicked.connect(self._on_skip_forward)
        controls_layout.addWidget(self.btn_skip_forward)

        self.btn_stop = QPushButton()
        self.btn_stop.setIcon(IconLoader.get_icon("stop", 24))
        self.btn_stop.setIconSize(QSize(24, 24))
        self.btn_stop.setMinimumSize(48, 48)
        self.btn_stop.setObjectName("white_button")
        self.btn_stop.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_stop.clicked.connect(self._on_stop)
        controls_layout.addWidget(self.btn_stop)

        controls_layout.addSpacing(24)

        # Vitesse
        speed_label = QLabel("Vitesse")
        speed_label.setObjectName("secondary_label")
        controls_layout.addWidget(speed_label)

        self.speed_combo = QComboBox()
        self.speed_combo.addItems(["0.5x", "0.75x", "1.0x", "1.25x", "1.5x", "2.0x"])
        self.speed_combo.setCurrentText("1.0x")
        self.speed_combo.currentTextChanged.connect(self._on_speed_changed)
        self.speed_combo.setMinimumWidth(80)
        controls_layout.addWidget(self.speed_combo)

        controls_layout.addSpacing(16)

        # Volume avec icône SVG
        volume_icon = QLabel()
        volume_icon.setPixmap(IconLoader.get_icon("mute", 20).pixmap(20, 20))
        controls_layout.addWidget(volume_icon)

        self.volume_slider = QSlider(Qt.Orientation.Horizontal)
        self.volume_slider.setRange(0, 100)
        self.volume_slider.setValue(100)
        self.volume_slider.setMaximumWidth(120)
        self.volume_slider.setCursor(Qt.CursorShape.PointingHandCursor)
        self.volume_slider.valueChanged.connect(self._on_volume_changed)
        controls_layout.addWidget(self.volume_slider)

        controls_layout.addStretch()

        layout.addLayout(controls_layout)

        return timeline

    def _setup_menu_bar(self):
        """Configure la barre de menu."""
        menubar = self.menuBar()

        # Menu Fichier
        file_menu = menubar.addMenu("Fichier")

        # Importer un fichier
        import_action = QAction("Importer un fichier audio", self)
        import_action.setShortcut(QKeySequence("Ctrl+O"))
        import_action.triggered.connect(self._on_import_file)
        file_menu.addAction(import_action)

        file_menu.addSeparator()

        # Transcrire
        transcribe_action = QAction("Transcrire", self)
        transcribe_action.setShortcut(QKeySequence("Ctrl+T"))
        transcribe_action.triggered.connect(self._start_transcription)
        file_menu.addAction(transcribe_action)

        file_menu.addSeparator()

        # Exporter TXT
        export_txt_action = QAction("Exporter au format TXT", self)
        export_txt_action.setShortcut(QKeySequence("Ctrl+S"))
        export_txt_action.triggered.connect(lambda: self._export_transcription("txt"))
        file_menu.addAction(export_txt_action)

        # Exporter DOCX
        export_docx_action = QAction("Exporter au format DOCX", self)
        export_docx_action.setShortcut(QKeySequence("Ctrl+Shift+S"))
        export_docx_action.triggered.connect(lambda: self._export_transcription("docx"))
        file_menu.addAction(export_docx_action)

        # Menu JuryAIssist
        app_menu = menubar.addMenu("JuryAIssist")

        # À propos
        about_action = QAction("À propos", self)
        about_action.triggered.connect(self._show_about)
        app_menu.addAction(about_action)

    def _apply_styles(self):
        """Applique les styles."""
        self.setStyleSheet(MAIN_STYLE)

    def _setup_shortcuts(self):
        """Configure les raccourcis clavier."""
        # Installer un event filter pour gérer les raccourcis globalement
        self.installEventFilter(self)

        print("✓ Raccourcis clavier configurés:")
        print("  Ctrl+O (Cmd+O sur Mac)     : Importer un fichier")
        print("  Ctrl+T (Cmd+T sur Mac)     : Transcrire")
        print("  Ctrl+S (Cmd+S sur Mac)     : Exporter TXT")
        print("  Ctrl+Shift+S               : Exporter DOCX")
        print("  Espace                     : Play/Pause")
        print("  ← Flèche gauche            : Reculer de 5s")
        print("  → Flèche droite            : Avancer de 5s")
        print("  ↑ Flèche haut              : Augmenter le volume")
        print("  ↓ Flèche bas               : Diminuer le volume")

    def eventFilter(self, obj, event):
        """Filtre les événements pour gérer les raccourcis clavier globalement."""
        if event.type() == QEvent.Type.KeyPress:
            key = event.key()

            # Ignorer les raccourcis si l'éditeur a le focus et qu'on tape du texte
            # Sauf pour les flèches et espace qu'on veut toujours intercepter
            if self.text_editor.hasFocus():
                # Si l'utilisateur modifie activement le texte (pas juste focus passif)
                # on laisse passer les touches normales
                if key not in [Qt.Key.Key_Left, Qt.Key.Key_Right, Qt.Key.Key_Up, Qt.Key.Key_Down, Qt.Key.Key_Space]:
                    return super().eventFilter(obj, event)

            # Gestion des raccourcis
            if key == Qt.Key.Key_Space:
                self._on_play()
                return True
            elif key == Qt.Key.Key_Left:
                self._on_skip_back()
                return True
            elif key == Qt.Key.Key_Right:
                self._on_skip_forward()
                return True
            elif key == Qt.Key.Key_Up:
                self._on_volume_up()
                return True
            elif key == Qt.Key.Key_Down:
                self._on_volume_down()
                return True

        return super().eventFilter(obj, event)

    # === SLOTS ===

    @Slot(float)
    def _on_position_changed(self, position: float):
        """Mise à jour de la position."""
        if not self.progress_slider.isSliderDown():
            duration = self.controller.get_duration()
            if duration > 0:
                progress = int((position / duration) * 1000)
                self.progress_slider.setValue(progress)

        self.time_label.setText(self._format_time(position))

    @Slot(float)
    def _on_duration_changed(self, duration: float):
        """Mise à jour de la durée."""
        self.duration_label.setText(self._format_time(duration))

    @Slot(object)
    def _on_state_changed(self, state: PlayerState):
        """Mise à jour de l'état."""
        if state == PlayerState.PLAYING:
            # Note: pas d'icône pause, on garde play pour l'instant
            self.btn_play.setIcon(IconLoader.get_icon("play", 28))
            self.statusBar().showMessage("Lecture en cours...")
        else:
            self.btn_play.setIcon(IconLoader.get_icon("play", 28))
            if state == PlayerState.STOPPED:
                self.statusBar().showMessage("Arrêté")
            else:
                self.statusBar().showMessage("En pause")

    @Slot(str)
    def _on_source_loaded(self, source_name: str):
        """Fichier chargé."""
        self.file_title.setText(source_name)
        self.statusBar().showMessage(f"Fichier chargé: {source_name}")

    @Slot()
    def _on_pedal_connected(self):
        """Pédale connectée."""
        self.pedal_badge.setText("● Pédale: Connectée (RS-31)")
        self.pedal_badge.setProperty("connected", "true")
        self.pedal_badge.setStyleSheet("")  # Force refresh
        self.pedal_badge.style().unpolish(self.pedal_badge)
        self.pedal_badge.style().polish(self.pedal_badge)

    @Slot()
    def _on_pedal_disconnected(self):
        """Pédale déconnectée."""
        self.pedal_badge.setText("○ Pédale: Non détectée")
        self.pedal_badge.setProperty("connected", "false")
        self.pedal_badge.setStyleSheet("")
        self.pedal_badge.style().unpolish(self.pedal_badge)
        self.pedal_badge.style().polish(self.pedal_badge)

    @Slot(object)
    def _on_pedal_action(self, action):
        """Action pédale."""
        from src.devices.pedal import PedalAction
        if action == PedalAction.PLAY_PAUSE:
            self.controller.toggle_play_pause()
        elif action == PedalAction.REWIND:
            self.controller.skip_backward(self.skip_duration)
        elif action == PedalAction.FORWARD:
            self.controller.skip_forward(self.skip_duration)

    # === HANDLERS ===

    def _on_import_file(self):
        """Importe un fichier audio."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Sélectionner un fichier audio",
            "",
            "Fichiers audio (*.mp3 *.wav *.m4a *.flac *.ogg);;Tous (*)"
        )

        if file_path:
            path = Path(file_path)
            item = TranscriptionItem(path, path.stem)
            self.transcriptions.append(item)

            # Ajouter à la liste avec icône SVG
            list_item = QListWidgetItem(item.name)
            list_item.setIcon(IconLoader.get_icon("radio", 16))
            self.transcriptions_list.addItem(list_item)
            self.transcriptions_list.setCurrentItem(list_item)

            # Charger le fichier
            self.current_transcription = item
            self.controller.load_file(file_path)

            # Proposer transcription
            reply = QMessageBox.question(
                self,
                "Transcription",
                "Voulez-vous transcrire ce fichier maintenant?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                self._start_transcription()

    def _on_transcription_selected(self, item: QListWidgetItem):
        """Sélection d'une transcription."""
        index = self.transcriptions_list.row(item)
        if 0 <= index < len(self.transcriptions):
            trans = self.transcriptions[index]
            self.current_transcription = trans
            self.text_editor.setText(trans.transcription_text)
            self.controller.load_file(str(trans.file_path))

    def _on_play(self):
        """Play/Pause."""
        self.controller.toggle_play_pause()

    def _on_stop(self):
        """Stop."""
        self.controller.stop()

    def _on_skip_back(self):
        """Reculer."""
        self.controller.skip_backward(self.skip_duration)

    def _on_skip_forward(self):
        """Avancer."""
        self.controller.skip_forward(self.skip_duration)

    def _on_slider_moved(self, value: int):
        """Slider déplacé."""
        duration = self.controller.get_duration()
        if duration > 0:
            position = (value / 1000.0) * duration
            self.controller.seek(position)

    def _on_speed_changed(self, text: str):
        """Vitesse changée."""
        speed = float(text.replace("x", ""))
        self.controller.set_speed(speed)

    def _on_volume_changed(self, value: int):
        """Volume changé."""
        self.controller.set_volume(value)

    def _on_volume_up(self):
        """Augmente le volume de 10%."""
        current = self.volume_slider.value()
        new_value = min(100, current + 10)
        self.volume_slider.setValue(new_value)

    def _on_volume_down(self):
        """Diminue le volume de 10%."""
        current = self.volume_slider.value()
        new_value = max(0, current - 10)
        self.volume_slider.setValue(new_value)

    def _on_text_changed(self):
        """Texte modifié."""
        if not self.is_modified and self.current_transcription:
            self.is_modified = True
            self.file_title.setText(f"{self.current_transcription.name} *")

    def _on_show_info(self):
        """Affiche les infos du fichier."""
        if not self.current_transcription:
            QMessageBox.information(self, "Information", "Aucun fichier chargé")
            return

        path = self.current_transcription.file_path
        duration = self.controller.get_duration()
        size_mb = path.stat().st_size / (1024 * 1024)

        info = f"""
<b>Nom:</b> {path.name}<br>
<b>Chemin:</b> {path}<br>
<b>Taille:</b> {size_mb:.1f} MB<br>
<b>Durée:</b> {self._format_time(duration)}<br>
<b>Format:</b> {path.suffix.upper()}
        """
        QMessageBox.information(self, "Informations du fichier", info)

    def _show_about(self):
        """Affiche les informations À propos."""
        about_text = """
<h2>JuryAIssist</h2>
<p><b>Version:</b> 1.0.0</p>
<p><b>Description:</b> Application de transcription audio intelligente pour juristes</p>
<br>
<p><b>Fonctionnalités:</b></p>
<ul>
<li>Transcription automatique avec Whisper AI</li>
<li>Lecture audio avec contrôle pédale RS-31</li>
<li>Export en TXT et DOCX</li>
<li>Interface moderne et intuitive</li>
</ul>
<br>
<p><b>Technologies:</b></p>
<ul>
<li>PySide6 (Qt for Python)</li>
<li>OpenAI Whisper</li>
<li>VLC Media Player</li>
</ul>
<br>
<p>© 2026 JuryAIssist - Tous droits réservés</p>
        """
        QMessageBox.about(self, "À propos de JuryAIssist", about_text)

    def _on_settings(self):
        """Affiche les paramètres."""
        dialog = SettingsDialog(self)
        dialog.model_combo.setCurrentText(self.whisper_model)
        dialog.skip_spin.setValue(int(self.skip_duration))

        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.whisper_model = dialog.model_combo.currentText()
            self.whisper_language = dialog.get_language_code()
            self.skip_duration = float(dialog.skip_spin.value())
            self.statusBar().showMessage("Paramètres sauvegardés", 3000)

    def _start_transcription(self):
        """Lance la transcription."""
        if not self.current_transcription:
            return

        # Dialogue de progression moderne
        progress = ModernProgressDialog(
            "Transcription",
            "Transcription en cours avec Whisper AI...",
            self
        )
        progress.set_status(f"Modèle: {self.whisper_model} | Langue: {self.whisper_language}")
        progress.show()

        # Forcer le rafraîchissement
        QApplication.processEvents()

        # Créer signaux pour communication thread-safe
        signals = TranscriptionSignals()
        signals.completed.connect(self._on_transcription_complete)
        signals.error.connect(self._on_transcription_error)

        def transcribe():
            try:
                print(f"Début transcription: {self.current_transcription.file_path}")  # Debug
                print(f"Modèle: {self.whisper_model}, Langue: {self.whisper_language}")  # Debug

                # Créer un transcriber avec le modèle sélectionné
                transcriber = WhisperTranscriber(model_size=self.whisper_model)
                result = transcriber.transcribe(
                    str(self.current_transcription.file_path),
                    language=self.whisper_language
                )

                print(f"Transcription terminée, status: {result.status}")  # Debug
                # Émettre signal thread-safe
                signals.completed.emit(result, progress)
            except Exception as e:
                print(f"Erreur transcription: {e}")  # Debug
                import traceback
                traceback.print_exc()
                signals.error.emit(str(e), progress)

        # Démarrer dans un thread séparé (non-daemon pour garantir completion)
        thread = threading.Thread(target=transcribe, daemon=False)
        thread.start()

    def _on_transcription_complete(self, result, progress):
        """Transcription terminée."""
        print("=== _on_transcription_complete appelé ===")  # Debug
        print(f"Result status: {result.status}")  # Debug
        print(f"Progress dialog visible: {progress.isVisible() if progress else 'None'}")  # Debug

        # result est un TranscriptionResult, pas un dict
        from src.transcription.transcriber import TranscriptionStatus

        # Vérifier le status avant de fermer le dialogue
        if result.status != TranscriptionStatus.COMPLETED:
            # Fermer le dialogue d'abord
            if progress:
                progress.close()
                progress.deleteLater()
            QApplication.processEvents()

            QMessageBox.critical(self, "Erreur", f"La transcription a échoué: {result.error_message}")
            return

        print(f"Segments reçus: {len(result.segments)}")  # Debug
        print(f"Texte complet: {len(result.full_text)} caractères")  # Debug

        text = result.full_text
        segments = result.segments

        # Sauvegarder dans l'objet transcription
        if self.current_transcription:
            self.current_transcription.transcription_text = text
            self.current_transcription.segments = segments
            print(f"Transcription sauvegardée dans {self.current_transcription.name}")  # Debug

        # Afficher avec timestamps
        formatted_text = ""
        for seg in segments:
            timestamp = self._format_time(seg.start)
            formatted_text += f"[{timestamp}] {seg.text.strip()}\n\n"

        print(f"Texte formaté: {len(formatted_text)} caractères")  # Debug

        # Bloquer les signaux pour éviter de déclencher is_modified
        self.text_editor.blockSignals(True)
        self.text_editor.setPlainText(formatted_text)
        self.text_editor.blockSignals(False)
        self.is_modified = False

        print("Texte inséré dans l'éditeur")  # Debug

        # Forcer la fermeture du dialogue APRÈS avoir mis à jour l'UI
        if progress:
            progress.close()
            progress.deleteLater()
            print("Dialogue fermé")  # Debug

        # Forcer le rafraîchissement complet de l'UI
        self.text_editor.update()
        self.text_editor.viewport().update()
        QApplication.processEvents()

        print("Affichage du message de succès")  # Debug
        # Afficher le message de succès
        QMessageBox.information(
            self,
            "Succès",
            f"Transcription terminée!\n\n{len(segments)} segments | {result.word_count} mots"
        )

        print("=== _on_transcription_complete terminé ===")  # Debug

    def _on_transcription_error(self, error, progress):
        """Erreur transcription."""
        progress.close()
        QMessageBox.critical(self, "Erreur", f"Erreur:\n{error}")

    def _format_time(self, seconds: float) -> str:
        """Formate le temps."""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes}:{secs:02d}"

    def _export_transcription(self, format_type: str = "txt"):
        """Exporte la transcription."""
        if not self.current_transcription or not self.current_transcription.transcription_text:
            QMessageBox.information(self, "Information", "Aucune transcription à exporter")
            return

        # Dialogue de sauvegarde
        if format_type == "txt":
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Exporter en TXT",
                f"{self.current_transcription.name}_transcription.txt",
                "Fichiers texte (*.txt)"
            )

            if file_path:
                try:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(self.text_editor.toPlainText())
                    QMessageBox.information(self, "Succès", f"Transcription exportée:\n{file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "Erreur", f"Erreur d'export:\n{e}")

        elif format_type == "docx":
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Exporter en DOCX",
                f"{self.current_transcription.name}_transcription.docx",
                "Documents Word (*.docx)"
            )

            if file_path:
                try:
                    from docx import Document
                    doc = Document()
                    doc.add_heading("Transcription Audio", 0)
                    doc.add_paragraph(f"Fichier: {self.current_transcription.name}")
                    doc.add_paragraph("")

                    # Ajouter le texte
                    for line in self.text_editor.toPlainText().split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)

                    doc.save(file_path)
                    QMessageBox.information(self, "Succès", f"Transcription exportée:\n{file_path}")
                except ImportError:
                    QMessageBox.critical(self, "Erreur", "Le module 'python-docx' n'est pas installé.\n\nInstallez-le avec: pip install python-docx")
                except Exception as e:
                    QMessageBox.critical(self, "Erreur", f"Erreur d'export:\n{e}")

    def closeEvent(self, event):
        """Fermeture."""
        if self.is_modified:
            reply = QMessageBox.question(
                self,
                "Sauvegarder?",
                "Modifications non sauvegardées. Quitter quand même?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                event.ignore()
                return

        self.controller.release()
        self.pedal.disconnect()
        event.accept()
