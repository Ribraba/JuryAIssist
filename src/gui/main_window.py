"""
Fenêtre principale de l'application.

Layout:
+--------------------------------------------------------+
| [Sidebar] | [Main Content]                            |
|           | ┌─────────────────────────────────────┐   |
|           | | Scrolling Transcript Timeline       |   |
|           | └─────────────────────────────────────┘   |
|           | ┌──────────────┬──────────────────────┐   |
|           | | Transcript   | Editor               |   |
|           | | (read-only)  | (editable)           |   |
|           | └──────────────┴──────────────────────┘   |
|           | ┌─────────────────────────────────────┐   |
|           | | Audio Controls                      |   |
|           | └─────────────────────────────────────┘   |
|           |             [Pedal Badge]                 |
+--------------------------------------------------------+

Principes SOLID:
- Single Responsibility: Gère l'assemblage des composants et la coordination
- Open/Closed: Extensible via ajout de nouvelles fonctionnalités
- Liskov Substitution: Interface compatible avec QMainWindow
- Interface Segregation: Signaux et slots bien définis
- Dependency Inversion: Injecte les dépendances (controller, transcriber, etc.)
- Tell, Don't Ask: Commande les widgets plutôt que d'interroger leur état
"""
from pathlib import Path
from typing import Optional, List

from PyQt5.QtCore import Qt, QTimer, QThread, pyqtSignal
from PyQt5.QtWidgets import (
    QMainWindow,
    QWidget,
    QHBoxLayout,
    QVBoxLayout,
    QSplitter,
    QFileDialog,
    QMessageBox,
    QProgressDialog,
    QDialog,
    QLabel,
)
from PyQt5.QtGui import QKeySequence

from src.audio.controller import AudioController
from src.audio.player import PlayerState
from src.audio.vlc_player import VLCAudioPlayer
from src.audio.timeline import TimeUtils

# Import des widgets
from src.gui.widgets.sidebar import SidebarWidget
from src.gui.widgets.scrolling_transcript_timeline import ScrollingTranscriptTimeline
from src.gui.widgets.editor import EditorPanel
from src.gui.widgets.audio_controls_panel import AudioControlsPanel
from src.gui.widgets.pedal_status_badge import PedalStatusBadge

# Pédale
from src.devices.olympus_pedal import OlympusPedal
from src.devices.pedal import PedalAction

# Transcription
from src.transcription.whisper_transcriber import WhisperTranscriber
from src.transcription.transcriber import TranscriptionSegment, TranscriptionStatus
from src.transcription.word_sync import WordSynchronizer

# Styles et ressources
from src.gui.theme import get_app_stylesheet, AppSpacing
from src.gui.resources import load_fonts, get_font

# Workers
from src.gui.workers import TranscriptionWorker

# Configuration
from src.config import get_settings

# Dialogs
from src.gui.dialogs import SettingsDialog, AboutDialog


class MainWindow(QMainWindow):
    """
    Fenêtre principale selon le design Figma.

    Nouvelle interface utilisateur moderne et épurée.
    """

    def __init__(self):
        """Initialise la fenêtre."""
        super().__init__()

        # Charger les polices personnalisées
        load_fonts()

        # Charger les paramètres
        self.settings = get_settings()

        # État
        self._current_audio_file: Optional[str] = None
        self._current_transcript: List[TranscriptionSegment] = []

        # Créer le player et le contrôleur
        self._player = VLCAudioPlayer()
        self._controller = AudioController(self._player)

        # WordSynchronizer pour la navigation par clic sur mot
        self._word_sync = WordSynchronizer()

        # Pédale Olympus (optionnelle)
        self._pedal: Optional[OlympusPedal] = None
        self._pedal_connected = False

        # Timer pour détecter la pédale périodiquement
        self._pedal_detect_timer = QTimer()
        self._pedal_detect_timer.timeout.connect(self._check_pedal_connection)
        self._pedal_detect_timer.setInterval(2000)  # Vérifier toutes les 2 secondes

        # Configuration de la fenêtre
        self.setWindowTitle("JuryAIssist - Transcription Audio Juridique")
        self.setMinimumSize(1200, 800)

        # Restaurer la taille depuis les paramètres
        window_width = self.settings.get("window_width", 1440)
        window_height = self.settings.get("window_height", 960)
        self.resize(window_width, window_height)

        # Appliquer le style global (selon le paramètre dark_mode)
        dark_mode = self.settings.get("dark_mode", False)
        if dark_mode:
            from src.gui.theme import get_app_stylesheet_dark
            self.setStyleSheet(get_app_stylesheet_dark())
        else:
            self.setStyleSheet(get_app_stylesheet())

        # Créer l'interface
        self._create_ui()
        self._create_menu()
        self._create_shortcuts()

        # Connecter les signaux
        self._connect_signals()

        # Restaurer les paramètres audio
        self._restore_audio_settings()

        # Tenter de connecter la pédale au démarrage
        self._init_pedal()

        # Démarrer le timer de détection de pédale
        self._pedal_detect_timer.start()

    def _restore_audio_settings(self):
        """Restaure les paramètres audio depuis les settings."""
        # Restaurer le volume
        volume = self.settings.get("volume", 70)
        self._controller.set_volume(volume)
        self._audio_controls.set_volume(volume)

        # Restaurer la vitesse
        speed = self.settings.get("playback_speed", 1.0)
        self._controller.set_speed(speed)
        self._audio_controls.set_speed(speed)

    def closeEvent(self, event):
        """Sauvegarde les paramètres avant de fermer."""
        # Sauvegarder la taille de la fenêtre
        self.settings.set("window_width", self.width())
        self.settings.set("window_height", self.height())

        # Sauvegarder le volume et la vitesse actuels
        self.settings.set("volume", self._controller.get_volume())
        self.settings.set("playback_speed", self._controller.get_speed())

        # Sauvegarder sur disque
        self.settings.save()

        # Nettoyer
        if self._transcription_worker and self._transcription_worker.isRunning():
            self._transcription_worker.stop()
            self._transcription_worker.wait()

        if self._controller:
            self._controller.release()

        if self._pedal:
            self._pedal.disconnect()

        event.accept()

    def _create_ui(self):
        """Crée l'interface utilisateur."""
        # Widget central
        central = QWidget()
        self.setCentralWidget(central)

        # Layout principal (horizontal: Sidebar | Content)
        main_layout = QHBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # 1. Sidebar (gauche)
        self._sidebar = SidebarWidget()
        main_layout.addWidget(self._sidebar)

        # 2. Zone de contenu principale (droite)
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)
        content_layout.setContentsMargins(AppSpacing.MD, AppSpacing.MD, AppSpacing.MD, AppSpacing.MD)
        content_layout.setSpacing(AppSpacing.MD)

        # 2a. Badge de pédale (en haut à droite)
        pedal_container = QWidget()
        pedal_layout = QHBoxLayout(pedal_container)
        pedal_layout.setContentsMargins(0, 0, 0, 0)
        pedal_layout.addStretch()
        self._pedal_badge = PedalStatusBadge()
        pedal_layout.addWidget(self._pedal_badge)
        content_layout.addWidget(pedal_container)

        # 2b. Panneau d'édition (PREND TOUT L'ESPACE)
        self._editor_panel = EditorPanel()
        content_layout.addWidget(self._editor_panel, 1)  # Prend tout l'espace disponible

        # 2c. Timeline avec transcription défilante cliquable (EN BAS, au-dessus des contrôles audio)
        self._timeline = ScrollingTranscriptTimeline()
        content_layout.addWidget(self._timeline)

        # 2d. Étiquette "Lecture" sous la timeline
        lecture_label = QLabel("Lecture")
        lecture_label.setObjectName("lectureLabel")
        lecture_label.setFont(get_font(16, 400))
        lecture_label.setStyleSheet("color: #444444; padding-left: 0px;")
        content_layout.addWidget(lecture_label)

        # 2d. Contrôles audio (tout en bas)
        self._audio_controls = AudioControlsPanel()
        content_layout.addWidget(self._audio_controls)

        main_layout.addWidget(content_widget, 1)

        # NE PAS initialiser de fichiers statiques - la liste commence vide

    def _create_menu(self):
        """Crée la barre de menu."""
        from PyQt5.QtWidgets import QAction, QMenu

        menubar = self.menuBar()

        # Menu Fichier
        file_menu = menubar.addMenu("&Fichier")

        # Action: Ouvrir
        open_action = QAction("&Ouvrir un fichier audio...", self)
        open_action.setShortcut(QKeySequence.Open)
        open_action.triggered.connect(self._load_file)
        file_menu.addAction(open_action)

        file_menu.addSeparator()

        # Action: Exporter TXT
        export_txt_action = QAction("Exporter en &TXT...", self)
        export_txt_action.triggered.connect(self._export_txt)
        file_menu.addAction(export_txt_action)

        # Action: Exporter DOCX
        export_docx_action = QAction("Exporter en &DOCX...", self)
        export_docx_action.triggered.connect(self._export_docx)
        file_menu.addAction(export_docx_action)

        file_menu.addSeparator()

        # Action: Quitter
        quit_action = QAction("&Quitter", self)
        quit_action.setShortcut(QKeySequence.Quit)
        quit_action.triggered.connect(self.close)
        file_menu.addAction(quit_action)

        # Menu Edition
        edit_menu = menubar.addMenu("&Edition")

        # Action: Copier
        copy_action = QAction("&Copier", self)
        copy_action.setShortcut(QKeySequence.Copy)
        edit_menu.addAction(copy_action)

        # Menu Transcription
        transcription_menu = menubar.addMenu("&Transcription")

        # Action: Transcrire le fichier audio
        transcribe_action = QAction("&Transcrire le fichier audio...", self)
        transcribe_action.setShortcut("Ctrl+T")
        transcribe_action.triggered.connect(self._start_transcription)
        transcription_menu.addAction(transcribe_action)

        # Menu Aide
        help_menu = menubar.addMenu("&Aide")

        # Action: À propos
        about_action = QAction("À &propos", self)
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)

    def _create_shortcuts(self):
        """Crée les raccourcis clavier."""
        from PyQt5.QtWidgets import QAction

        # Space = Play/Pause
        play_pause_shortcut = QAction(self)
        play_pause_shortcut.setShortcut(Qt.Key_Space)
        play_pause_shortcut.triggered.connect(self._controller.toggle_play_pause)
        self.addAction(play_pause_shortcut)

        # Flèche droite = Avancer de 5 secondes
        skip_forward_shortcut = QAction(self)
        skip_forward_shortcut.setShortcut(Qt.Key_Right)
        skip_forward_shortcut.triggered.connect(lambda: self._controller.skip_forward(5.0))
        self.addAction(skip_forward_shortcut)

        # Flèche gauche = Reculer de 5 secondes
        skip_backward_shortcut = QAction(self)
        skip_backward_shortcut.setShortcut(Qt.Key_Left)
        skip_backward_shortcut.triggered.connect(lambda: self._controller.skip_backward(5.0))
        self.addAction(skip_backward_shortcut)

    def _connect_signals(self):
        """Connecte tous les signaux."""
        # === Sidebar ===
        self._sidebar.file_selected.connect(self._on_file_selected)
        self._sidebar.import_clicked.connect(self._load_file)
        self._sidebar.settings_clicked.connect(self._show_settings)
        # search_changed: à implémenter plus tard

        # === Timeline ===
        self._timeline.position_clicked.connect(self._controller.seek)

        # === Audio Controls ===
        self._audio_controls.play_clicked.connect(self._controller.play)
        self._audio_controls.pause_clicked.connect(self._controller.pause)
        self._audio_controls.stop_clicked.connect(self._controller.stop)
        self._audio_controls.skip_forward_clicked.connect(
            lambda: self._controller.skip_forward(5.0)
        )
        self._audio_controls.skip_backward_clicked.connect(
            lambda: self._controller.skip_backward(5.0)
        )
        self._audio_controls.speed_changed.connect(self._controller.set_speed)
        self._audio_controls.volume_changed.connect(self._on_volume_changed)

        # === Audio Controller → UI ===
        self._controller.position_changed.connect(self._on_position_changed)
        self._controller.duration_changed.connect(self._on_duration_changed)
        self._controller.state_changed.connect(self._on_state_changed)

        # === Timeline (transcription cliquable) → Audio ===
        # TEMPORAIREMENT DÉSACTIVÉ: peut causer des crashs
        # TODO: Réactiver quand la synchronisation mot-timestamp sera stable
        # self._timeline.word_clicked.connect(self._on_word_clicked)

    def _load_file(self):
        """Charge un fichier audio."""
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Ouvrir un fichier audio",
            str(Path.home()),
            "Audio (*.mp3 *.wav *.m4a *.flac *.ogg *.dss);;Tous (*)",
        )

        if path:
            if self._controller.load_file(path):
                self._current_audio_file = path
                filename = Path(path).name

                # Mettre à jour l'éditeur avec le nom du fichier
                self._editor_panel.set_filename(filename)

                # Activer les contrôles
                self._audio_controls.enable_controls(True)

                # Remplacer le fichier actuel dans la sidebar (un seul import visible)
                self._sidebar.clear_transcript_files()
                self._sidebar.add_transcript_file(filename, selected=True)

                # Lancer la transcription automatiquement ? Ou via menu ?
                reply = QMessageBox.question(
                    self,
                    "Transcription",
                    f"Fichier chargé: {filename}\n\nVoulez-vous le transcrire maintenant ?",
                    QMessageBox.Yes | QMessageBox.No
                )

                if reply == QMessageBox.Yes:
                    self._start_transcription()

            else:
                QMessageBox.warning(
                    self,
                    "Erreur",
                    f"Impossible de charger le fichier:\n{path}"
                )

    def _start_transcription(self):
        """Lance la transcription du fichier audio."""
        if not self._current_audio_file:
            QMessageBox.warning(
                self,
                "Attention",
                "Veuillez d'abord charger un fichier audio."
            )
            return

        # Utiliser les paramètres sauvegardés
        model_size = self.settings.get("preferred_model", "base")
        language = self.settings.get("preferred_language", "fr")

        # Créer le worker de transcription
        self._transcription_worker = TranscriptionWorker(
            self._current_audio_file,
            model_size,
            language
        )

        # Créer la barre de progression
        self._progress_dialog = QProgressDialog(
            "Transcription en cours...",
            "Annuler",
            0,
            100,
            self
        )
        self._progress_dialog.setWindowModality(Qt.WindowModal)
        self._progress_dialog.setMinimumDuration(0)
        self._progress_dialog.canceled.connect(self._transcription_worker.stop)

        # Connecter les signaux
        self._transcription_worker.progress.connect(self._progress_dialog.setValue)
        self._transcription_worker.finished.connect(self._on_transcription_finished)
        self._transcription_worker.error.connect(self._on_transcription_error)

        # Lancer la transcription
        self._transcription_worker.start()

    def _on_transcription_finished(self, segments: List[TranscriptionSegment]):
        """Appelé quand la transcription est terminée."""
        self._progress_dialog.close()
        self._current_transcript = segments

        # Construire l'index de synchronisation mot-timestamp
        self._word_sync.build_index(segments)

        # Afficher dans la timeline (transcription cliquable)
        self._timeline.set_transcript_segments(segments)

        # Afficher dans l'éditeur
        text_lines = []
        for seg in segments:
            text_lines.append(seg.text)
        self._editor_panel.set_text("\n\n".join(text_lines))

        QMessageBox.information(
            self,
            "Succès",
            f"Transcription terminée avec succès !\n\n"
            f"{len(segments)} segments créés."
        )

    def _on_transcription_error(self, error_message: str):
        """Appelé en cas d'erreur de transcription."""
        self._progress_dialog.close()
        QMessageBox.critical(
            self,
            "Erreur",
            f"Erreur lors de la transcription :\n\n{error_message}"
        )

    def _on_file_selected(self, filename: str):
        """
        Appelé quand un fichier est sélectionné dans la sidebar.

        Args:
            filename: Nom du fichier
        """
        # TODO: Charger le fichier sélectionné
        self._editor_panel.set_filename(filename)

    def _show_settings(self):
        """Affiche les paramètres."""
        dialog = SettingsDialog(self)
        dialog.settings_changed.connect(self._on_settings_changed)
        dialog.exec_()

    def _on_settings_changed(self):
        """Appelé quand les paramètres sont modifiés."""
        # Recharger les paramètres
        self.settings = get_settings()

        # Appliquer le nouveau thème si le mode sombre a changé
        dark_mode = self.settings.get("dark_mode", False)
        if dark_mode:
            from src.gui.theme import get_app_stylesheet_dark
            self.setStyleSheet(get_app_stylesheet_dark())
        else:
            self.setStyleSheet(get_app_stylesheet())

        # Reconfigurer la pédale si elle est connectée
        if self._pedal and self._pedal_connected:
            # Créer un nouveau mapper avec les nouveaux paramètres
            new_mapper = self._create_pedal_mapper()
            # Mettre à jour le mapper de la pédale
            # Note: ButtonActionMapper n'expose pas de méthode pour changer le mapper
            # après initialisation, donc on devra reconnecter la pédale
            # Pour l'instant, informer l'utilisateur
            needs_restart = True
        else:
            needs_restart = False

        # Informer l'utilisateur
        message = "Les paramètres ont été sauvegardés.\n\n"
        if dark_mode or needs_restart:
            message += "Note: Redémarrez l'application pour appliquer tous les changements."

        QMessageBox.information(
            self,
            "Paramètres sauvegardés",
            message
        )

    def _export_txt(self):
        """Exporte la transcription en TXT."""
        text = self._editor_panel.get_text()
        if not text:
            QMessageBox.warning(self, "Attention", "Aucun texte à exporter.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporter en TXT",
            str(Path.home() / "transcription.txt"),
            "Fichiers texte (*.txt)"
        )

        if path:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(text)
                QMessageBox.information(self, "Succès", f"Exporté: {Path(path).name}")
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Erreur lors de l'export:\n{e}")

    def _export_docx(self):
        """Exporte la transcription en DOCX."""
        text = self._editor_panel.get_text()
        if not text:
            QMessageBox.warning(self, "Attention", "Aucun texte à exporter.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporter en DOCX",
            str(Path.home() / "transcription.docx"),
            "Documents Word (*.docx)"
        )

        if path:
            try:
                from docx import Document

                doc = Document()
                doc.add_heading('Transcription Audio', 0)

                # Ajouter le texte édité (paragraphe par paragraphe)
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():  # Ignorer lignes vides
                        doc.add_paragraph(para)

                doc.save(path)
                QMessageBox.information(self, "Succès", f"Exporté en DOCX: {Path(path).name}")
            except ImportError:
                QMessageBox.warning(
                    self,
                    "Module manquant",
                    "Le module python-docx n'est pas installé.\n"
                    "Installez-le avec: pip install python-docx"
                )
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Erreur lors de l'export DOCX:\n{e}")

    def _show_about(self):
        """Affiche la boîte À propos."""
        QMessageBox.about(
            self,
            "À propos de JuryAIssist",
            "<h3>JuryAIssist</h3>"
            "<p>Application d'assistance à la transcription et à l'édition "
            "d'enregistrements audio pour juristes.</p>"
            "<p><b>Version:</b> 2.0</p>"
            "<p><b>Interface:</b> Moderne et intuitive</p>"
            "<p>Développé avec PyQt5 et les principes SOLID</p>"
        )

    def _on_position_changed(self, position: float):
        """Appelé quand la position change."""
        # Mettre à jour la timeline
        self._timeline.set_position(position)

        # Mettre à jour les contrôles audio
        duration = self._controller.get_duration()
        self._audio_controls.set_time(position, duration)

    def _on_duration_changed(self, duration: float):
        """Appelé quand la durée change."""
        self._timeline.set_duration(duration)

    def _on_state_changed(self, state: PlayerState):
        """Appelé quand l'état de lecture change."""
        is_playing = (state == PlayerState.PLAYING)
        self._audio_controls.set_playing_state(is_playing)

    def _on_volume_changed(self, volume: int):
        """
        Appelé quand le volume change.

        Args:
            volume: Volume (0-100)
        """
        # Passer directement au contrôleur (0-100)
        self._controller.set_volume(volume)

    def _on_word_clicked(self, word: str, timestamp: float):
        """
        Appelé quand l'utilisateur clique sur un mot dans la timeline.

        Args:
            word: Mot cliqué
            timestamp: Timestamp du mot en secondes (émis par ScrollingTranscriptTimeline)
        """
        if not self._current_transcript:
            return

        # Le timestamp est directement fourni par la timeline, pas besoin de chercher
        # Faire le seek vers ce timestamp
        self._controller.seek(timestamp)

    # === Pédale ===

    def _init_pedal(self):
        """Détecte et connecte la pédale (silencieux si absente)."""
        try:
            # Créer un mapper personnalisé basé sur les paramètres
            from src.devices.action_mapper import ButtonActionMapper
            mapper = self._create_pedal_mapper()

            # Créer la pédale avec le mapper personnalisé
            self._pedal = OlympusPedal(mapper=mapper)
            if self._pedal.detect() and self._pedal.connect():
                self._pedal_connected = True
                self._connect_pedal_signals()
                self._pedal_badge.set_connected(True, "RS-31")
                print("✅ Pédale connectée avec succès")
        except ImportError:
            # hidapi non disponible
            pass
        except Exception as e:
            print(f"⚠️ Erreur pédale: {e}")

    def _create_pedal_mapper(self):
        """
        Crée un ActionMapper basé sur les paramètres sauvegardés.

        Returns:
            ButtonActionMapper configuré selon les préférences utilisateur
        """
        from src.devices.action_mapper import ButtonActionMapper

        # Charger les actions depuis les paramètres
        button_mapping = {}
        for button_num in range(1, 5):
            action_str = self.settings.get(f"pedal_button_{button_num}", None)
            if action_str:
                # Convertir la string en PedalAction
                try:
                    action = PedalAction(action_str)
                    button_mapping[button_num] = action
                except ValueError:
                    # Valeur invalide, utiliser le défaut
                    pass

        # Créer le mapper (avec defaults si mapping vide)
        if button_mapping:
            return ButtonActionMapper(initial_mapping=button_mapping)
        else:
            return ButtonActionMapper()  # Utilise les defaults

    def _check_pedal_connection(self):
        """Vérifie périodiquement si une pédale est connectée."""
        # Si déjà connectée, vérifier qu'elle est toujours là
        if self._pedal_connected and self._pedal:
            if not self._pedal.is_connected():
                print("⚠️ Pédale déconnectée")
                self._pedal_connected = False
                self._pedal = None
                self._pedal_badge.set_connected(False)

        # Si pas connectée, essayer de détecter
        if not self._pedal_connected:
            self._init_pedal()

    def _connect_pedal_signals(self):
        """Connecte les signaux de la pédale."""
        if self._pedal:
            self._pedal.action_triggered.connect(self._on_pedal_action)

    def _on_pedal_action(self, action: PedalAction):
        """
        Gère les actions de la pédale.

        Le mapping action → commande se fait ici en fonction des paramètres.
        """
        if action == PedalAction.PLAY_PAUSE:
            self._controller.toggle_play_pause()
        elif action == PedalAction.SKIP_FORWARD:
            self._controller.skip_forward(5.0)
        elif action == PedalAction.SKIP_BACKWARD:
            self._controller.skip_backward(5.0)
        elif action == PedalAction.STOP:
            self._controller.stop()
        elif action == PedalAction.CYCLE_SPEED:
            self._controller.cycle_speed()

    def closeEvent(self, event):
        """Fermeture propre."""
        # Arrêter le timer
        self._pedal_detect_timer.stop()

        # Déconnecter la pédale
        if self._pedal:
            self._pedal.disconnect()

        # Libérer le controller
        self._controller.release()

        event.accept()
