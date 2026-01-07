"""
Fenêtre principale selon Phase 3.1.1 de la roadmap.

Layout :
+------------------------------------------+
|  Menu: Fichier | Edition | Aide         |
+------------------------------------------+
|  [Timeline avec curseur de lecture]      |
|  [Play] [Pause] [<<] [>>] [Vitesse]     |
+------------------------------------------+
|  Transcription  |  Éditeur               |
|  Brute          |  (modifiable)          |
|  (lecture seule)|                        |
+------------------------------------------+
|  Status: Position 00:02:35 / 01:23:45    |
+------------------------------------------+

Fonctionnalités :
- Splitter horizontal entre les deux panneaux
- Menu avec Fichier (Ouvrir, Exporter, Quitter)
- Barre de status avec position actuelle
- Raccourcis clavier (Space=Play/Pause, etc.)
"""

from pathlib import Path
from typing import Optional, List

from PyQt5.QtCore import Qt, QTimer, QThread, pyqtSignal
from PyQt5.QtWidgets import (
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QSplitter,
    QFileDialog,
    QStatusBar,
    QAction,
    QMessageBox,
    QProgressDialog,
    QDialog,
    QVBoxLayout as QVBoxLayoutDialog,
    QLabel,
    QComboBox,
    QPushButton,
    QHBoxLayout,
)
from PyQt5.QtGui import QKeySequence

from src.audio.controller import AudioController
from src.audio.player import PlayerState
from src.audio.vlc_player import VLCAudioPlayer
from src.audio.timeline import TimeUtils

# Import des widgets
from src.gui.widgets.timeline_widget import TimelineWidget
from src.gui.widgets.audio_controls import AudioControlsWidget
from src.gui.widgets.editor_panel import EditorPanel
from src.gui.widgets.transcript_panel import TranscriptPanel

# Pédale
from src.devices.olympus_pedal import OlympusPedal
from src.devices.pedal import PedalAction

# Transcription
from src.transcription.whisper_transcriber import WhisperTranscriber
from src.transcription.transcriber import TranscriptionSegment, TranscriptionStatus
from src.transcription.word_sync import WordSynchronizer

# Styles
from src.gui.styles import get_app_style


class MainWindow(QMainWindow):
    """
    Fenêtre principale selon la roadmap Phase 3.

    Fournit une interface complète pour la transcription audio.
    """

    def __init__(self):
        """Initialise la fenêtre."""
        super().__init__()

        # État
        self._current_audio_file: Optional[str] = None
        self._current_transcript = []

        # Créer le player et le contrôleur
        self._player = VLCAudioPlayer()
        self._controller = AudioController(self._player)

        # WordSynchronizer pour la navigation par clic sur mot
        self._word_sync = WordSynchronizer()

        # Pédale Olympus (optionnelle)
        self._pedal: Optional[OlympusPedal] = None
        self._pedal_connected = False

        # Timer pour détecter la pédale périodiquement
        self._pedal_detect_timer = QTimer()
        self._pedal_detect_timer.timeout.connect(self._check_pedal_connection)
        self._pedal_detect_timer.setInterval(2000)  # Vérifier toutes les 2 secondes

        # Configuration de la fenêtre
        self.setWindowTitle("JuryAIssist - Transcription Audio Juridique")
        self.setMinimumSize(1000, 700)
        self.resize(1200, 800)

        # Appliquer le style global
        self.setStyleSheet(get_app_style())

        # Créer l'interface
        self._create_ui()
        self._create_menu()
        self._create_status_bar()
        self._create_shortcuts()

        # Connecter les signaux
        self._connect_signals()

        # Tenter de connecter la pédale au démarrage
        self._init_pedal()

        # Démarrer le timer de détection de pédale
        self._pedal_detect_timer.start()

    def _create_ui(self):
        """Crée l'interface utilisateur selon la roadmap."""
        # Widget central
        central = QWidget()
        self.setCentralWidget(central)

        # Layout principal
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(16, 16, 16, 16)
        main_layout.setSpacing(12)

        # Timeline
        self._timeline = TimelineWidget()
        main_layout.addWidget(self._timeline)

        # Contrôles audio
        self._audio_controls = AudioControlsWidget()
        main_layout.addWidget(self._audio_controls)

        # Splitter horizontal (Transcription Brute | Éditeur)
        splitter = QSplitter(Qt.Horizontal)

        # Panneau de transcription brute (gauche)
        self._transcript_panel = TranscriptPanel()
        splitter.addWidget(self._transcript_panel)

        # Panneau d'édition (droite)
        self._editor_panel = EditorPanel()
        splitter.addWidget(self._editor_panel)

        # Ratio 50/50
        splitter.setSizes([500, 500])

        main_layout.addWidget(splitter, 1)  # Prend tout l'espace disponible

    def _create_menu(self):
        """Crée le menu selon la roadmap."""
        menubar = self.menuBar()

        # Menu Fichier
        file_menu = menubar.addMenu("&Fichier")

        # Action: Ouvrir
        open_action = QAction("&Ouvrir un fichier audio...", self)
        open_action.setShortcut(QKeySequence.Open)
        open_action.triggered.connect(self._load_file)
        file_menu.addAction(open_action)

        file_menu.addSeparator()

        # Action: Exporter TXT
        export_txt_action = QAction("Exporter en &TXT...", self)
        export_txt_action.triggered.connect(self._export_txt)
        file_menu.addAction(export_txt_action)

        # Action: Exporter DOCX
        export_docx_action = QAction("Exporter en &DOCX...", self)
        export_docx_action.triggered.connect(self._export_docx)
        file_menu.addAction(export_docx_action)

        file_menu.addSeparator()

        # Action: Quitter
        quit_action = QAction("&Quitter", self)
        quit_action.setShortcut(QKeySequence.Quit)
        quit_action.triggered.connect(self.close)
        file_menu.addAction(quit_action)

        # Menu Edition
        edit_menu = menubar.addMenu("&Edition")

        # Action: Copier
        copy_action = QAction("&Copier", self)
        copy_action.setShortcut(QKeySequence.Copy)
        edit_menu.addAction(copy_action)

        # Menu Transcription
        transcription_menu = menubar.addMenu("&Transcription")

        # Action: Transcrire le fichier audio
        transcribe_action = QAction("&Transcrire le fichier audio...", self)
        transcribe_action.setShortcut("Ctrl+T")
        transcribe_action.triggered.connect(self._start_transcription)
        transcription_menu.addAction(transcribe_action)

        # Menu Aide
        help_menu = menubar.addMenu("&Aide")

        # Action: À propos
        about_action = QAction("À &propos", self)
        about_action.triggered.connect(self._show_about)
        help_menu.addAction(about_action)

    def _create_status_bar(self):
        """Crée la barre de status selon la roadmap."""
        self._status_bar = QStatusBar()
        self.setStatusBar(self._status_bar)

        # Créer l'indicateur de pédale (label permanent à droite)
        from PyQt5.QtWidgets import QLabel
        self._pedal_status_label = QLabel()
        self._pedal_status_label.setStyleSheet("""
            QLabel {
                color: #999999;
                padding: 2px 8px;
                border-radius: 3px;
                font-size: 11px;
            }
        """)
        self._update_pedal_status_label(False)
        self._status_bar.addPermanentWidget(self._pedal_status_label)

        # Message par défaut
        self._update_status("Prêt")

    def _create_shortcuts(self):
        """Crée les raccourcis clavier selon la roadmap."""
        # Space = Play/Pause
        play_pause_shortcut = QAction(self)
        play_pause_shortcut.setShortcut(Qt.Key_Space)
        play_pause_shortcut.triggered.connect(self._controller.toggle_play_pause)
        self.addAction(play_pause_shortcut)

        # Ctrl+O = Ouvrir
        # Déjà défini dans le menu

    def _connect_signals(self):
        """Connecte tous les signaux."""
        # Timeline → Audio
        self._timeline.position_changed.connect(self._controller.seek)

        # Audio Controls → Audio
        self._audio_controls.play_clicked.connect(self._controller.play)
        self._audio_controls.pause_clicked.connect(self._controller.pause)
        self._audio_controls.stop_clicked.connect(self._controller.stop)
        self._audio_controls.skip_forward_clicked.connect(
            lambda: self._controller.skip_forward(5.0)
        )
        self._audio_controls.skip_backward_clicked.connect(
            lambda: self._controller.skip_backward(5.0)
        )
        self._audio_controls.speed_changed.connect(self._controller.set_speed)
        self._audio_controls.volume_changed.connect(self._controller.set_volume)

        # Audio → Timeline
        self._controller.position_changed.connect(self._timeline.set_position)
        self._controller.duration_changed.connect(self._timeline.set_duration)

        # Audio → Audio Controls
        self._controller.state_changed.connect(self._on_state_changed)

        # Audio → Status Bar
        self._controller.position_changed.connect(self._on_position_changed)

        # Transcript Panel → Audio (clic sur mot → seek)
        self._transcript_panel.word_clicked.connect(self._on_word_clicked)

        # Audio → Transcript Panel (surlignage en temps réel)
        self._controller.position_changed.connect(self._transcript_panel.highlight_segment_at_time)

    def _load_file(self):
        """Charge un fichier audio."""
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Ouvrir un fichier audio",
            str(Path.home()),
            "Audio (*.mp3 *.wav *.m4a *.flac *.ogg *.dss);;Tous (*)",
        )

        if path:
            if self._controller.load_file(path):
                self._current_audio_file = path
                self._audio_controls.enable_controls(True)
                self._update_status(f"Fichier chargé: {Path(path).name}")
            else:
                QMessageBox.warning(
                    self,
                    "Erreur",
                    f"Impossible de charger le fichier:\n{path}"
                )

    def _export_txt(self):
        """Exporte la transcription en TXT."""
        text = self._editor_panel.get_text()
        if not text:
            QMessageBox.warning(self, "Attention", "Aucun texte à exporter.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporter en TXT",
            str(Path.home() / "transcription.txt"),
            "Fichiers texte (*.txt)"
        )

        if path:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(text)
                self._update_status(f"Exporté: {Path(path).name}")
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Erreur lors de l'export:\n{e}")

    def _export_docx(self):
        """Exporte la transcription en DOCX."""
        text = self._editor_panel.get_text()
        if not text:
            QMessageBox.warning(self, "Attention", "Aucun texte à exporter.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporter en DOCX",
            str(Path.home() / "transcription.docx"),
            "Documents Word (*.docx)"
        )

        if path:
            try:
                from docx import Document

                doc = Document()
                doc.add_heading('Transcription Audio', 0)

                # Ajouter le texte édité (paragraphe par paragraphe)
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():  # Ignorer lignes vides
                        doc.add_paragraph(para)

                doc.save(path)
                self._update_status(f"Exporté en DOCX: {Path(path).name}")
                QMessageBox.information(self, "Succès", "Export DOCX réussi !")
            except ImportError:
                QMessageBox.warning(
                    self,
                    "Module manquant",
                    "Le module python-docx n'est pas installé.\n"
                    "Installez-le avec: pip install python-docx"
                )
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Erreur lors de l'export DOCX:\n{e}")

    def _start_transcription(self):
        """Lance la transcription du fichier audio."""
        if not self._current_audio_file:
            QMessageBox.warning(
                self,
                "Attention",
                "Veuillez d'abord charger un fichier audio."
            )
            return

        # Afficher une boîte de dialogue pour choisir le modèle
        dialog = TranscriptionDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            model_size = dialog.get_model_size()
            language = dialog.get_language()

            # Créer le worker de transcription
            self._transcription_worker = TranscriptionWorker(
                self._current_audio_file,
                model_size,
                language
            )

            # Créer la barre de progression
            self._progress_dialog = QProgressDialog(
                "Transcription en cours...",
                "Annuler",
                0,
                100,
                self
            )
            self._progress_dialog.setWindowModality(Qt.WindowModal)
            self._progress_dialog.setMinimumDuration(0)
            self._progress_dialog.canceled.connect(self._transcription_worker.stop)

            # Connecter les signaux
            self._transcription_worker.progress.connect(self._progress_dialog.setValue)
            self._transcription_worker.finished.connect(self._on_transcription_finished)
            self._transcription_worker.error.connect(self._on_transcription_error)

            # Lancer la transcription
            self._transcription_worker.start()
            self._update_status("Transcription en cours...")

    def _on_transcription_finished(self, segments: List[TranscriptionSegment]):
        """Appelé quand la transcription est terminée."""
        self._progress_dialog.close()
        self._current_transcript = segments

        # Construire l'index de synchronisation mot-timestamp
        self._word_sync.build_index(segments)

        # Afficher dans le panneau de transcription brute
        self._transcript_panel.set_transcript_segments(segments)

        # Afficher dans l'éditeur
        text_lines = []
        for seg in segments:
            text_lines.append(seg.text)
        self._editor_panel.set_text("\n\n".join(text_lines))

        self._update_status(f"Transcription terminée : {len(segments)} segments")
        QMessageBox.information(
            self,
            "Succès",
            f"Transcription terminée avec succès !\n\n"
            f"{len(segments)} segments créés."
        )

    def _on_transcription_error(self, error_message: str):
        """Appelé en cas d'erreur de transcription."""
        self._progress_dialog.close()
        self._update_status("Erreur de transcription")
        QMessageBox.critical(
            self,
            "Erreur",
            f"Erreur lors de la transcription :\n\n{error_message}"
        )

    def _on_word_clicked(self, word: str, position: int):
        """
        Appelé quand l'utilisateur clique sur un mot dans la transcription brute.

        Cherche le timestamp du mot et fait un seek vers cette position.

        Args:
            word: Mot cliqué
            position: Position du mot dans le texte
        """
        if not self._current_transcript:
            return

        # Chercher le mot à la position exacte cliquée
        word_timestamp = self._word_sync.find_word_at_position(position)

        if word_timestamp is not None:
            # Faire le seek vers ce timestamp
            self._controller.seek(word_timestamp.start_time)
            self._update_status(f"Navigation vers '{word}' à {word_timestamp.start_time:.1f}s")
        else:
            # Le mot n'a pas été trouvé dans l'index
            # Cela peut arriver si le mot a été édité par l'utilisateur
            pass

    def _show_about(self):
        """Affiche la boîte À propos."""
        QMessageBox.about(
            self,
            "À propos de JuryAIssist",
            "JuryAIssist v2.0.0\n\n"
            "Application de transcription audio juridique\n"
            "avec contrôle par pédale USB.\n\n"
            "© 2026"
        )

    def _on_state_changed(self, state: PlayerState):
        """Appelé quand l'état de lecture change."""
        is_playing = (state == PlayerState.PLAYING)
        self._audio_controls.set_playing_state(is_playing)

    def _on_position_changed(self, position: float):
        """Appelé quand la position change."""
        duration = self._controller.get_duration()
        if duration > 0:
            position_str = TimeUtils.seconds_to_timestamp(position)
            duration_str = TimeUtils.seconds_to_timestamp(duration)
            self._update_status(f"Position: {position_str} / {duration_str}")

            # Surligner le segment actuel dans la transcription brute
            self._transcript_panel.highlight_segment_at_time(position)

    def _update_status(self, message: str):
        """Met à jour la barre de status."""
        self._status_bar.showMessage(message)

    def _update_pedal_status_label(self, connected: bool):
        """
        Met à jour l'indicateur de statut de la pédale.

        Args:
            connected: True si la pédale est connectée
        """
        if connected:
            self._pedal_status_label.setText("🎮 Pédale connectée")
            self._pedal_status_label.setStyleSheet("""
                QLabel {
                    color: #30D158;
                    background-color: rgba(48, 209, 88, 0.15);
                    padding: 4px 12px;
                    border-radius: 6px;
                    font-size: 11px;
                    font-weight: 600;
                }
            """)
        else:
            self._pedal_status_label.setText("🎮 Pédale déconnectée")
            self._pedal_status_label.setStyleSheet("""
                QLabel {
                    color: #8E8E93;
                    padding: 4px 12px;
                    border-radius: 6px;
                    font-size: 11px;
                }
            """)

    # === Pédale ===

    def _init_pedal(self):
        """Détecte et connecte la pédale (silencieux si absente)."""
        try:
            self._pedal = OlympusPedal()
            if self._pedal.detect() and self._pedal.connect():
                self._pedal_connected = True
                self._connect_pedal_signals()
                self._update_pedal_status_label(True)
                print("✅ Pédale connectée avec succès")
        except ImportError:
            # hidapi non disponible
            pass
        except Exception as e:
            print(f"⚠️ Erreur pédale: {e}")

    def _check_pedal_connection(self):
        """Vérifie périodiquement si une pédale est connectée."""
        # Si déjà connectée, vérifier qu'elle est toujours là
        if self._pedal_connected and self._pedal:
            if not self._pedal.is_connected():
                print("⚠️ Pédale déconnectée")
                self._pedal_connected = False
                self._pedal = None
                self._update_pedal_status_label(False)

        # Si pas connectée, essayer de détecter
        if not self._pedal_connected:
            self._init_pedal()

    def _connect_pedal_signals(self):
        """Connecte les signaux de la pédale."""
        if self._pedal:
            self._pedal.action_triggered.connect(self._on_pedal_action)

    def _on_pedal_action(self, action: PedalAction):
        """Gère les actions de la pédale."""
        if action == PedalAction.PLAY_PAUSE:
            self._controller.toggle_play_pause()
        elif action == PedalAction.SKIP_FORWARD:
            self._controller.skip_forward(5.0)
        elif action == PedalAction.SKIP_BACKWARD:
            self._controller.skip_backward(5.0)
        elif action == PedalAction.STOP:
            self._controller.stop()
        elif action == PedalAction.CYCLE_SPEED:
            self._controller.cycle_speed()

    def closeEvent(self, event):
        """Fermeture propre."""
        # Arrêter le timer
        self._pedal_detect_timer.stop()

        # Déconnecter la pédale
        if self._pedal:
            self._pedal.disconnect()

        # Libérer le controller
        self._controller.release()

        event.accept()


# === Classes auxiliaires ===

class TranscriptionDialog(QDialog):
    """
    Boîte de dialogue pour configurer la transcription.
    """

    def __init__(self, parent=None):
        """Initialise la boîte de dialogue."""
        super().__init__(parent)

        self.setWindowTitle("Configuration de la transcription")
        self.setMinimumWidth(400)

        layout = QVBoxLayoutDialog()

        # Label d'information
        info_label = QLabel(
            "Veuillez choisir le modèle Whisper et la langue pour la transcription."
        )
        info_label.setWordWrap(True)
        layout.addWidget(info_label)

        # Choix du modèle
        model_layout = QHBoxLayout()
        model_label = QLabel("Modèle :")
        model_layout.addWidget(model_label)

        self._model_combo = QComboBox()
        self._model_combo.addItems(["tiny", "base", "small", "medium", "large"])
        self._model_combo.setCurrentText("base")
        model_layout.addWidget(self._model_combo)
        layout.addLayout(model_layout)

        # Choix de la langue
        lang_layout = QHBoxLayout()
        lang_label = QLabel("Langue :")
        lang_layout.addWidget(lang_label)

        self._lang_combo = QComboBox()
        self._lang_combo.addItems([
            "fr", "en", "es", "de", "it", "pt", "nl", "pl",
            "ru", "zh", "ja", "ko", "ar", "hi"
        ])
        self._lang_combo.setCurrentText("fr")
        lang_layout.addWidget(self._lang_combo)
        layout.addLayout(lang_layout)

        # Boutons
        button_layout = QHBoxLayout()
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        cancel_button = QPushButton("Annuler")
        cancel_button.clicked.connect(self.reject)
        button_layout.addWidget(cancel_button)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

        self.setLayout(layout)

    def get_model_size(self) -> str:
        """Retourne le modèle choisi."""
        return self._model_combo.currentText()

    def get_language(self) -> str:
        """Retourne la langue choisie."""
        return self._lang_combo.currentText()


class TranscriptionWorker(QThread):
    """
    Worker pour effectuer la transcription dans un thread séparé.
    """

    # Signaux
    progress = pyqtSignal(int)  # Pourcentage de progression
    finished = pyqtSignal(list)  # Liste de TranscriptSegment
    error = pyqtSignal(str)  # Message d'erreur

    def __init__(self, audio_file: str, model_size: str, language: str):
        """
        Initialise le worker.

        Args:
            audio_file: Chemin du fichier audio
            model_size: Taille du modèle Whisper
            language: Code de la langue
        """
        super().__init__()
        self._audio_file = audio_file
        self._model_size = model_size
        self._language = language
        self._should_stop = False

    def run(self):
        """Effectue la transcription."""
        try:
            # Créer le transcriber avec la taille de modèle choisie
            self.progress.emit(10)

            transcriber = WhisperTranscriber(model_size=self._model_size)
            self.progress.emit(30)

            if self._should_stop:
                return

            # Transcrire (le modèle sera chargé automatiquement)
            result = transcriber.transcribe(
                self._audio_file,
                language=self._language
            )

            self.progress.emit(90)

            if self._should_stop:
                return

            # Vérifier le résultat
            if result.status == TranscriptionStatus.ERROR:
                self.error.emit(result.error_message or "Erreur inconnue")
            else:
                self.progress.emit(100)
                self.finished.emit(result.segments)

        except Exception as e:
            self.error.emit(str(e))

    def stop(self):
        """Arrête la transcription."""
        self._should_stop = True
