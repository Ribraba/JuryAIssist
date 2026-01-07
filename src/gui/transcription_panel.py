"""
Panneau de transcription pour l'interface graphique.

Affiche les segments transcrits avec timing et permet l'export.
"""

from pathlib import Path
from typing import Optional

from PyQt5.QtCore import Qt, pyqtSignal, QObject, QThread
from PyQt5.QtWidgets import (
    QFrame,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QLabel,
    QTextEdit,
    QComboBox,
    QProgressBar,
    QFileDialog,
    QMessageBox,
)

from src.transcription.transcriber import TranscriptionResult, TranscriptionStatus
from src.transcription.whisper_transcriber import WhisperTranscriber
from src.gui.styles import (
    get_card_style,
    get_secondary_button_style,
    get_load_button_style,
    COLORS,
)


class TranscriptionWorker(QObject):
    """Worker thread pour transcription asynchrone."""

    finished = pyqtSignal(object)  # TranscriptionResult
    progress = pyqtSignal(str)  # Message de progression
    error = pyqtSignal(str)  # Message d'erreur

    def __init__(self, audio_path: str, model_size: str = "base", language: str = "fr"):
        super().__init__()
        self._audio_path = audio_path
        self._model_size = model_size
        self._language = language

    def run(self):
        """Exécute la transcription."""
        try:
            self.progress.emit("Chargement du modèle Whisper...")
            transcriber = WhisperTranscriber(model_size=self._model_size)

            self.progress.emit("Transcription en cours...")
            result = transcriber.transcribe(self._audio_path, language=self._language)

            transcriber.release()
            self.finished.emit(result)

        except Exception as e:
            self.error.emit(f"Erreur: {str(e)}")


class TranscriptionPanel(QFrame):
    """
    Panneau d'affichage et contrôle de la transcription.

    Permet de lancer une transcription et d'afficher les résultats.
    """

    def __init__(self):
        super().__init__()
        self.setObjectName("card")
        self.setStyleSheet(get_card_style())

        self._current_result: Optional[TranscriptionResult] = None
        self._thread: Optional[QThread] = None
        self._worker: Optional[TranscriptionWorker] = None

        self._create_ui()

    def _create_ui(self):
        """Crée l'interface."""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)

        # Titre
        title = QLabel("Transcription")
        title.setStyleSheet(f"""
            QLabel {{
                color: {COLORS['text_primary']};
                font-size: 20px;
                font-weight: 700;
            }}
        """)
        layout.addWidget(title)

        # Contrôles
        controls = self._create_controls()
        layout.addLayout(controls)

        # Barre de progression
        self._progress_bar = QProgressBar()
        self._progress_bar.setStyleSheet(f"""
            QProgressBar {{
                border: 1px solid {COLORS['border']};
                border-radius: 6px;
                text-align: center;
                background: {COLORS['bg_tertiary']};
                height: 20px;
            }}
            QProgressBar::chunk {{
                background: {COLORS['text_primary']};
                border-radius: 5px;
            }}
        """)
        self._progress_bar.setVisible(False)
        layout.addWidget(self._progress_bar)

        # Message de statut
        self._status_label = QLabel("")
        self._status_label.setStyleSheet(f"""
            QLabel {{
                color: {COLORS['text_secondary']};
                font-size: 13px;
            }}
        """)
        self._status_label.setVisible(False)
        layout.addWidget(self._status_label)

        # Zone de texte
        self._text_edit = QTextEdit()
        self._text_edit.setReadOnly(True)
        self._text_edit.setPlaceholderText("Le texte transcrit apparaîtra ici...")
        self._text_edit.setStyleSheet(f"""
            QTextEdit {{
                border: 1px solid {COLORS['border']};
                border-radius: 8px;
                padding: 12px;
                background: {COLORS['bg_secondary']};
                color: {COLORS['text_primary']};
                font-size: 14px;
                line-height: 1.6;
            }}
        """)
        layout.addWidget(self._text_edit, 1)

        # Boutons d'export
        export_layout = self._create_export_buttons()
        layout.addLayout(export_layout)

    def _create_controls(self) -> QHBoxLayout:
        """Crée les contrôles de transcription."""
        layout = QHBoxLayout()
        layout.setSpacing(12)

        # Sélection du modèle
        model_label = QLabel("Modèle:")
        model_label.setStyleSheet(f"color: {COLORS['text_secondary']};")
        layout.addWidget(model_label)

        self._model_combo = QComboBox()
        self._model_combo.addItems(["tiny", "base", "small", "medium"])
        self._model_combo.setCurrentText("base")
        self._model_combo.setStyleSheet(get_secondary_button_style())
        layout.addWidget(self._model_combo)

        layout.addStretch()

        # Bouton transcrire
        self._transcribe_btn = QPushButton("Transcrire")
        self._transcribe_btn.setObjectName("load")
        self._transcribe_btn.setStyleSheet(get_load_button_style())
        self._transcribe_btn.setCursor(Qt.PointingHandCursor)
        self._transcribe_btn.clicked.connect(self._start_transcription)
        self._transcribe_btn.setEnabled(False)
        layout.addWidget(self._transcribe_btn)

        return layout

    def _create_export_buttons(self) -> QHBoxLayout:
        """Crée les boutons d'export."""
        layout = QHBoxLayout()
        layout.setSpacing(12)

        layout.addStretch()

        # Export TXT
        self._export_txt_btn = QPushButton("Exporter TXT")
        self._export_txt_btn.setObjectName("secondary")
        self._export_txt_btn.setStyleSheet(get_secondary_button_style())
        self._export_txt_btn.setCursor(Qt.PointingHandCursor)
        self._export_txt_btn.clicked.connect(self._export_txt)
        self._export_txt_btn.setEnabled(False)
        layout.addWidget(self._export_txt_btn)

        # Export DOCX
        self._export_docx_btn = QPushButton("Exporter DOCX")
        self._export_docx_btn.setObjectName("secondary")
        self._export_docx_btn.setStyleSheet(get_secondary_button_style())
        self._export_docx_btn.setCursor(Qt.PointingHandCursor)
        self._export_docx_btn.clicked.connect(self._export_docx)
        self._export_docx_btn.setEnabled(False)
        layout.addWidget(self._export_docx_btn)

        return layout

    def set_audio_file(self, audio_path: Optional[str]):
        """Définit le fichier audio à transcrire."""
        self._audio_path = audio_path
        self._transcribe_btn.setEnabled(audio_path is not None)

    def _start_transcription(self):
        """Lance la transcription."""
        if not hasattr(self, '_audio_path') or not self._audio_path:
            return

        # Désactiver les boutons
        self._transcribe_btn.setEnabled(False)
        self._export_txt_btn.setEnabled(False)
        self._export_docx_btn.setEnabled(False)

        # Afficher la progression
        self._progress_bar.setVisible(True)
        self._progress_bar.setRange(0, 0)  # Mode indéterminé
        self._status_label.setVisible(True)
        self._status_label.setText("Initialisation...")

        # Créer le worker et le thread
        self._thread = QThread()
        self._worker = TranscriptionWorker(
            self._audio_path,
            model_size=self._model_combo.currentText(),
            language="fr"
        )
        self._worker.moveToThread(self._thread)

        # Connecter les signaux
        self._thread.started.connect(self._worker.run)
        self._worker.finished.connect(self._on_transcription_finished)
        self._worker.progress.connect(self._on_progress)
        self._worker.error.connect(self._on_error)
        self._worker.finished.connect(self._thread.quit)
        self._worker.finished.connect(self._worker.deleteLater)
        self._thread.finished.connect(self._thread.deleteLater)

        # Démarrer
        self._thread.start()

    def _on_progress(self, message: str):
        """Mise à jour de la progression."""
        self._status_label.setText(message)

    def _on_transcription_finished(self, result: TranscriptionResult):
        """Transcription terminée."""
        self._current_result = result

        # Cacher la progression
        self._progress_bar.setVisible(False)

        if result.status == TranscriptionStatus.COMPLETED:
            # Afficher le résultat
            self._display_result(result)
            self._status_label.setText(
                f"Transcription terminée • {len(result.segments)} segments • "
                f"{result.word_count} mots"
            )

            # Activer les boutons d'export
            self._export_txt_btn.setEnabled(True)
            self._export_docx_btn.setEnabled(True)
        else:
            self._status_label.setText(f"Erreur: {result.error_message}")

        # Réactiver le bouton transcrire
        self._transcribe_btn.setEnabled(True)

    def _on_error(self, error_message: str):
        """Erreur lors de la transcription."""
        self._progress_bar.setVisible(False)
        self._status_label.setText(error_message)
        self._transcribe_btn.setEnabled(True)

        QMessageBox.critical(self, "Erreur", error_message)

    def _display_result(self, result: TranscriptionResult):
        """Affiche le résultat de la transcription."""
        # Format avec timestamps
        text = ""
        for seg in result.segments:
            timestamp = f"[{seg.start:.1f}s - {seg.end:.1f}s]"
            text += f"{timestamp} {seg.text}\n\n"

        self._text_edit.setPlainText(text)

    def _export_txt(self):
        """Exporte en fichier texte."""
        if not self._current_result:
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporter en TXT",
            str(Path.home() / "transcription.txt"),
            "Fichiers texte (*.txt)"
        )

        if path:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    # Texte avec timestamps
                    for seg in self._current_result.segments:
                        f.write(f"[{seg.start:.1f}s - {seg.end:.1f}s] {seg.text}\n")

                QMessageBox.information(self, "Succès", "Export TXT réussi !")
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Erreur d'export: {str(e)}")

    def _export_docx(self):
        """Exporte en fichier Word."""
        if not self._current_result:
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporter en DOCX",
            str(Path.home() / "transcription.docx"),
            "Documents Word (*.docx)"
        )

        if path:
            try:
                from docx import Document

                doc = Document()
                doc.add_heading('Transcription Audio', 0)

                # Informations
                doc.add_paragraph(f"Langue: {self._current_result.language}")
                doc.add_paragraph(f"Segments: {len(self._current_result.segments)}")
                doc.add_paragraph(f"Durée: {self._current_result.duration:.1f}s")
                doc.add_heading('Texte', level=1)

                # Segments
                for seg in self._current_result.segments:
                    p = doc.add_paragraph()
                    p.add_run(f"[{seg.start:.1f}s - {seg.end:.1f}s] ").bold = True
                    p.add_run(seg.text)

                doc.save(path)
                QMessageBox.information(self, "Succès", "Export DOCX réussi !")
            except ImportError:
                QMessageBox.warning(
                    self, "Module manquant",
                    "Le module python-docx n'est pas installé.\n"
                    "Installez-le avec: pip install python-docx"
                )
            except Exception as e:
                QMessageBox.critical(self, "Erreur", f"Erreur d'export: {str(e)}")
